"""Word document (.docx) generation — mirrors proposal_pdf.html section by section."""
import io
import logging

logger = logging.getLogger(__name__)

_C_BLUE  = (0x17, 0x60, 0xD6)
_C_DARK  = (0x0C, 0x1C, 0x42)
_C_MUTED = (0x5A, 0x6D, 0x8C)
_C_LIGHT = (0xD0, 0xE1, 0xFA)
_C_BG    = (0xF5, 0xF9, 0xFF)
_C_WHITE = (0xFF, 0xFF, 0xFF)
_C_GREEN = (0x00, 0x9E, 0x73)
_C_EDF   = (0xED, 0xF4, 0xFF)


def _rgb(r, g, b):
    from docx.shared import RGBColor
    return RGBColor(r, g, b)


def _hex_fill(cell, rgb_tuple):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '%02X%02X%02X' % rgb_tuple)
    # Remove existing shd if any
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)


def _get_or_add_tblPr(tbl):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    return tblPr


def _no_borders(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = _get_or_add_tblPr(table._tbl)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _thin_borders(table, rgb=(0xD0, 0xE1, 0xFA)):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hex_color = '%02X%02X%02X' % rgb
    tblPr = _get_or_add_tblPr(table._tbl)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), hex_color)
        tblBorders.append(el)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _cell_para(cell, text, bold=False, size_pt=10, color=None, italic=False, align=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(str(text))
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if color:
        r.font.color.rgb = _rgb(*color)
    return p


def _add_cell_line(cell, text, bold=False, size_pt=10, color=None, italic=False, align=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.add_paragraph()
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(str(text))
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if color:
        r.font.color.rgb = _rgb(*color)
    return p


def generate_proposal_docx(proposal) -> bytes | None:
    """Generate Word document matching proposal_pdf.html structure. Returns bytes or None."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error('python-docx not installed')
        return None

    try:
        doc = Document()

        # ── Page setup ───────────────────────────────────────────────────────────
        for sec in doc.sections:
            sec.top_margin    = Cm(1.8)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin   = Cm(2.0)
            sec.right_margin  = Cm(2.0)

        # Default style
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.font.color.rgb = _rgb(*_C_DARK)

        c = proposal.company_data or {}
        issued = (proposal.issued_at.strftime('%d/%m/%Y')
                  if hasattr(proposal.issued_at, 'strftime') else str(proposal.issued_at or ''))

        # ═══════════════════════════════════════════════════════════════════════
        # HEADER — mirrors PDF: logo/company left (blue), meta right (dark)
        # ═══════════════════════════════════════════════════════════════════════
        hdr = doc.add_table(rows=1, cols=2)
        hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_borders(hdr)
        hdr.columns[0].width = Cm(11)
        hdr.columns[1].width = Cm(7)

        lh = hdr.rows[0].cells[0]
        rh = hdr.rows[0].cells[1]
        _hex_fill(lh, _C_BLUE)
        _hex_fill(rh, _C_DARK)
        lh.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        rh.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Left: company name
        _cell_para(lh, c.get('trade_name', 'Web-Impulsa'),
                   bold=True, size_pt=22, color=_C_WHITE)
        _add_cell_line(lh, c.get('website', 'webimpulsa.es'),
                       size_pt=9, color=(0xBF, 0xDB, 0xFE))

        # Right: proposal metadata
        _cell_para(rh, 'PROPUESTA COMERCIAL',
                   bold=True, size_pt=9, color=_C_WHITE)
        _add_cell_line(rh, f'Nº {proposal.number}',
                       bold=False, size_pt=10, color=_C_WHITE)
        _add_cell_line(rh, f'Fecha: {issued}',
                       size_pt=9, color=(0xBF, 0xDB, 0xFE))
        _add_cell_line(rh, f'Válido: {proposal.valid_days} días naturales',
                       size_pt=9, color=(0xBF, 0xDB, 0xFE))

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # EMISOR / CLIENTE — 2-column info box
        # ═══════════════════════════════════════════════════════════════════════
        info = doc.add_table(rows=1, cols=2)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER
        _thin_borders(info)
        info.columns[0].width = Cm(9)
        info.columns[1].width = Cm(9)

        li = info.rows[0].cells[0]
        ri = info.rows[0].cells[1]
        _hex_fill(li, _C_BG)
        _hex_fill(ri, _C_BG)
        li.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        ri.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Emisor
        _cell_para(li, 'EMISOR', bold=True, size_pt=8, color=_C_BLUE)
        _add_cell_line(li, c.get('trade_name', 'Web-Impulsa'), bold=True, size_pt=11, color=_C_DARK)
        if c.get('legal_name'): _add_cell_line(li, c['legal_name'], size_pt=10, color=_C_DARK)
        if c.get('nif'):        _add_cell_line(li, f"NIF/NIE: {c['nif']}", size_pt=10, color=_C_DARK)
        if c.get('address'):    _add_cell_line(li, c['address'], size_pt=10, color=_C_DARK)
        if c.get('city'):       _add_cell_line(li, c['city'], size_pt=10, color=_C_DARK)
        if c.get('email'):      _add_cell_line(li, c['email'], size_pt=10, color=_C_DARK)
        if c.get('phone'):      _add_cell_line(li, c['phone'], size_pt=10, color=_C_DARK)

        # Cliente
        _cell_para(ri, 'PREPARADO PARA', bold=True, size_pt=8, color=_C_BLUE)
        _add_cell_line(ri, proposal.client_name or '—', bold=True, size_pt=11, color=_C_DARK)
        if proposal.client_email:    _add_cell_line(ri, proposal.client_email, size_pt=10, color=_C_DARK)
        if proposal.client_phone:    _add_cell_line(ri, proposal.client_phone, size_pt=10, color=_C_DARK)
        if proposal.client_biz_type: _add_cell_line(ri, proposal.client_biz_type, size_pt=10, color=_C_MUTED)
        _add_cell_line(ri, '', size_pt=6)
        _add_cell_line(ri, 'NIF / CIF (para factura):', size_pt=8, color=_C_MUTED)
        _add_cell_line(ri, '___________________________________', size_pt=10, color=_C_MUTED)
        _add_cell_line(ri, 'Dirección fiscal:', size_pt=8, color=_C_MUTED)
        _add_cell_line(ri, '___________________________________', size_pt=10, color=_C_MUTED)

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # PROJECT SUMMARY LINE
        # ═══════════════════════════════════════════════════════════════════════
        ps = doc.add_paragraph()
        ps.paragraph_format.space_after = Pt(2)
        r1 = ps.add_run('Proyecto: ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = _rgb(*_C_DARK)
        r2 = ps.add_run(proposal.project_name or proposal.package or '—')
        r2.font.size = Pt(10); r2.font.color.rgb = _rgb(*_C_MUTED)
        r3 = ps.add_run('   Plazo estimado: ')
        r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = _rgb(*_C_DARK)
        r4 = ps.add_run(proposal.timeline or '—')
        r4.font.size = Pt(10); r4.font.color.rgb = _rgb(*_C_MUTED)

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # PRICING TABLE — same rows as PDF
        # ═══════════════════════════════════════════════════════════════════════
        ph = doc.add_paragraph('DESGLOSE ECONÓMICO')
        ph.runs[0].bold = True; ph.runs[0].font.size = Pt(8)
        ph.runs[0].font.color.rgb = _rgb(*_C_BLUE)
        ph.paragraph_format.space_after = Pt(3)

        subtotal_bd = proposal.package_base_price + proposal.extras_price + proposal.rush_amount

        price_rows = [
            # (label, value, is_header, is_total, is_discount, is_muted)
            ('Concepto',         'Importe',                         True,  False, False, False),
            (proposal.package or 'Paquete base',
             f'{proposal.package_base_price} €',                    False, False, False, False),
        ]
        for e in (proposal.extras or []):
            price_rows.append((f"  + {e.get('name','')}",
                               f"{e.get('price',0)} €",            False, False, False, True))
        if proposal.rush and proposal.rush_amount:
            price_rows.append(('  + Urgencia (×1.25)',
                               f'+{proposal.rush_amount} €',        False, False, False, True))
        price_rows += [
            ('Subtotal',                f'{subtotal_bd} €',         False, False, False, True),
            (f'Descuento {proposal.discount_pct}%',
             f'−{proposal.discount_amount} €',                      False, False, True,  False),
            ('Base imponible',          f'{proposal.taxable_base} €', False, False, False, True),
            ('IVA 21%',                 f'{proposal.iva_amount} €', False, False, False, True),
            ('TOTAL CON IVA',           f'{proposal.total_with_iva} €', False, True, False, False),
        ]
        if proposal.maintenance_plan:
            price_rows.append((f'Mantenimiento: {proposal.maintenance_plan}',
                               f'{proposal.maintenance_price} €/mes', False, False, False, True))
        if proposal.hours_plan_name:
            price_rows.append((proposal.hours_plan_name,
                               f'{proposal.hours_plan_price} €/mes', False, False, False, True))
        if proposal.hosting_price:
            price_rows.append(('Dominio + Hosting',
                               f'{proposal.hosting_price} €/mes', False, False, False, True))

        pt = doc.add_table(rows=len(price_rows), cols=2)
        pt.alignment = WD_TABLE_ALIGNMENT.CENTER
        _thin_borders(pt)
        pt.columns[0].width = Cm(13)
        pt.columns[1].width = Cm(5)

        for i, (label, value, is_hdr, is_tot, is_disc, is_muted) in enumerate(price_rows):
            lc = pt.rows[i].cells[0]
            rc = pt.rows[i].cells[1]
            if is_hdr:
                _hex_fill(lc, _C_BLUE); _hex_fill(rc, _C_BLUE)
                col = _C_WHITE
            elif is_tot:
                _hex_fill(lc, _C_EDF); _hex_fill(rc, _C_EDF)
                col = _C_BLUE
            elif is_disc:
                col = _C_GREEN
            elif is_muted:
                col = _C_MUTED
            else:
                col = _C_DARK

            _cell_para(lc, label, bold=(is_hdr or is_tot), size_pt=10, color=col)
            _cell_para(rc, value, bold=(is_hdr or is_tot), size_pt=10, color=col, align='right')

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # ALCANCE / FUERA DE ALCANCE — 2 columns
        # ═══════════════════════════════════════════════════════════════════════
        sc = doc.add_table(rows=1, cols=2)
        sc.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_borders(sc)
        sc.columns[0].width = Cm(10)
        sc.columns[1].width = Cm(8)

        sl = sc.rows[0].cells[0]
        sr = sc.rows[0].cells[1]
        sl.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        sr.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        _cell_para(sl, 'ALCANCE DEL PROYECTO', bold=True, size_pt=8, color=_C_BLUE)
        for item in (proposal.scope or []):
            _add_cell_line(sl, f'› {item}', size_pt=9.5, color=_C_DARK)

        _cell_para(sr, 'FUERA DE ALCANCE', bold=True, size_pt=8, color=_C_BLUE)
        for item in (proposal.out_of_scope or []):
            _add_cell_line(sr, f'· {item}', size_pt=9, color=_C_MUTED)

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # FASES / PAGO — 2 columns
        # ═══════════════════════════════════════════════════════════════════════
        fp = doc.add_table(rows=1, cols=2)
        fp.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_borders(fp)
        fp.columns[0].width = Cm(10)
        fp.columns[1].width = Cm(8)

        fl = fp.rows[0].cells[0]
        fr = fp.rows[0].cells[1]
        fl.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        fr.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        _cell_para(fl, 'FASES DEL PROYECTO', bold=True, size_pt=8, color=_C_BLUE)
        for phase in (proposal.phases or []):
            _add_cell_line(fl, f'› {phase}', size_pt=9.5, color=_C_DARK)

        _cell_para(fr, 'FORMA DE PAGO', bold=True, size_pt=8, color=_C_BLUE)
        _hex_fill(fr, _C_EDF)
        _add_cell_line(fr, 'Pago en dos etapas', bold=True, size_pt=10, color=_C_DARK)
        half = round(proposal.taxable_base / 2)
        _add_cell_line(fr, '', size_pt=4)
        _add_cell_line(fr, '50% al inicio', bold=True, size_pt=10, color=_C_BLUE)
        _add_cell_line(fr, f'{half} € (sin IVA)', size_pt=9, color=_C_DARK)
        _add_cell_line(fr, '', size_pt=4)
        _add_cell_line(fr, '50% a la entrega', bold=True, size_pt=10, color=_C_BLUE)
        _add_cell_line(fr, f'{half} € (sin IVA)', size_pt=9, color=_C_DARK)
        _add_cell_line(fr, '', size_pt=4)
        _add_cell_line(fr, 'Entrega final tras el pago completo.',
                       size_pt=8, color=_C_MUTED, italic=True)

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # CONDICIONES GENERALES
        # ═══════════════════════════════════════════════════════════════════════
        ph2 = doc.add_paragraph('CONDICIONES GENERALES')
        ph2.runs[0].bold = True; ph2.runs[0].font.size = Pt(8)
        ph2.runs[0].font.color.rgb = _rgb(*_C_BLUE)
        ph2.paragraph_format.space_after = Pt(3)

        for i, cond in enumerate(proposal.conditions or [], 1):
            pc = doc.add_paragraph()
            pc.paragraph_format.space_before = Pt(0)
            pc.paragraph_format.space_after  = Pt(2)
            pc.paragraph_format.left_indent  = Cm(0.4)
            r = pc.add_run(f'{i}.  {cond}')
            r.font.size = Pt(8.5)
            r.font.color.rgb = _rgb(*_C_MUTED)

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # FIRMA — 4-column table, mirrors PDF signature block
        # ═══════════════════════════════════════════════════════════════════════
        ph3 = doc.add_paragraph('ACEPTACIÓN')
        ph3.runs[0].bold = True; ph3.runs[0].font.size = Pt(8)
        ph3.runs[0].font.color.rgb = _rgb(*_C_BLUE)
        ph3.paragraph_format.space_after = Pt(2)

        ps2 = doc.add_paragraph(
            'La firma de este documento implica aceptación del alcance, precio y condiciones indicadas.'
        )
        ps2.runs[0].font.size = Pt(8.5)
        ps2.runs[0].font.color.rgb = _rgb(*_C_MUTED)
        ps2.runs[0].italic = True
        ps2.paragraph_format.space_after = Pt(6)

        sig = doc.add_table(rows=2, cols=4)
        sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        _thin_borders(sig, (0xE2, 0xE8, 0xF0))
        for j, label in enumerate(['Nombre completo *', 'NIF / CIF', 'Fecha', 'Firma']):
            lc = sig.rows[0].cells[j]
            bc = sig.rows[1].cells[j]
            _cell_para(lc, label, bold=True, size_pt=8, color=_C_MUTED)
            _cell_para(bc, '', size_pt=14)
            _add_cell_line(bc, '', size_pt=14)

        doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════════
        # FOOTER
        # ═══════════════════════════════════════════════════════════════════════
        foot_text = (
            f"{c.get('trade_name','Web-Impulsa')} · {c.get('email','info@webimpulsa.es')}"
            f" · {c.get('phone','')} · {c.get('website','webimpulsa.es')}"
            + (f" · NIF {c['nif']}" if c.get('nif') else '')
        )
        pf = doc.add_paragraph(foot_text)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.runs[0].font.size = Pt(8)
        pf.runs[0].font.color.rgb = _rgb(*_C_MUTED)

        # ── Serialize ────────────────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = buf.read()
        logger.info('DOCX generated: proposal %s (%d bytes)', proposal.number, len(result))
        return result

    except Exception as exc:
        logger.error('DOCX generation failed for proposal %s: %s',
                     getattr(proposal, 'number', '?'), exc, exc_info=True)
        return None
